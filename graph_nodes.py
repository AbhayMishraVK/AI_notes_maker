import re
import time
from langchain_core.output_parsers import StrOutputParser, JsonOutputParser
from langchain_core.prompts import PromptTemplate
import ast
import os
import docx
from docx import Document
from agent_prompts.prompt_difference import prompt_for_difference
from agent_prompts.prompt_example import prompt_for_example
from agent_prompts.prompt_for_reason import prompt_for_reason
from agent_prompts.prompt_intro_defination import prompt_for_intro_defination
from agent_prompts.prompt_make_list import prompt_to_convert_list
from agent_prompts.prompt_advantage import prompt_advantage_disadvantage
from agent_prompts.prompt_real_life_example import prompt_for_real_life_example
from agent_prompts.prompt_transcript import prompt_for_transcript
from agent_prompts.prompt_types import prompt_for_types
from common_info import model_load
from youtube_video_transcript import youtube_video_transcript
from common_info import time_to_stop
from dotenv import load_dotenv

load_dotenv()

abhay_api_key = os.getenv("GRAQ_API_KEY")
asmit_api_key = os.getenv("ASMIT_GRAQ_API_KEY")
alfaiz_api_key = os.getenv("ALFAIZ_GRAQ_API_KEY")


def starting_node(state):
    return state


def take_url_and_make_summary_node(state):
    list_of_url = state['list_of_url']
    current_url = list_of_url[0]
    transcribe_text = youtube_video_transcript(url_of_video=current_url)
    prompt = PromptTemplate(
        template=prompt_for_transcript(transcribe_text),
        input_variables=[],
    )
    llm = model_load(abhay_api_key, temperature=1)
    make_summary = prompt | llm | StrOutputParser
    del list_of_url[0]
    state['summary_of_video'] = make_summary
    state['list_of_url'] = list_of_url
    return state


def take_question_and_make_list_node(state):
    user_question = state['user_question']
    create_list_of_question = re.split(r'\n+', user_question.strip())
    print("\n", create_list_of_question)
    state['list_of_question'] = create_list_of_question
    return state


def get_current_question_node(state):
    list_of_question = state['list_of_question']
    current_question = list_of_question[0]
    state['current_question'] = current_question
    del list_of_question[0]
    state['list_of_question'] = list_of_question
    return state


def create_intro_definition_node(state):
    youtube_video_present = state['youtube_video_present']
    summary_of_video = state['summary_of_video']
    current_topic = state['current_question']
    create_intro_defination = None
    llm = model_load(asmit_api_key, temperature=1)

    if current_topic is not None:
        prompt = PromptTemplate(
            template=prompt_for_intro_defination(topic_name=current_topic),
            input_variables=[],
        )
        create_intro_defination = prompt | llm | StrOutputParser()
    elif summary_of_video is not None:
        prompt = PromptTemplate(
            template=prompt_for_intro_defination(youtube_summary=summary_of_video),
            input_variables=[],
        )
        create_intro_defination = prompt | llm | StrOutputParser()

    create_intro_defination = create_intro_defination.invoke({})
    state['definition'] = create_intro_defination
    return state


def create_advantage_disadvantage_node(state):
    intro_and_defination = state['definition']
    prompt = PromptTemplate(
        template=prompt_advantage_disadvantage(intro_and_defination=intro_and_defination),
        input_variables=[],
    )
    llm = model_load(alfaiz_api_key, temperature=0.7)
    advantage_disadvantage = prompt | llm | StrOutputParser()
    advantage_disadvantage = advantage_disadvantage.invoke({})
    if advantage_disadvantage != 'Not possible':
        state['advantage_and_disadvantage'] = advantage_disadvantage
    return state


def create_difference_node(state):
    intro_and_defination = state['definition']
    prompt = PromptTemplate(
        template=prompt_for_difference(intro_and_defination),
        input_variables=[],
    )
    llm = model_load(abhay_api_key, temperature=0.7)
    difference = prompt | llm | StrOutputParser()
    difference = difference.invoke({})
    if difference != 'Not possible':
        state['difference'] = difference
    return state


def create_types_node(state):
    intro_and_defination = state['definition']
    prompt = PromptTemplate(
        template=prompt_for_types(intro_and_defination=intro_and_defination),
        input_variables=[],
    )
    llm = model_load(asmit_api_key, temperature=0.7)
    types = prompt | llm | StrOutputParser()
    types = types.invoke({})
    if types != 'Not possible':
        state['types'] = types
    return state


def create_reason_node(state):
    intro_and_defination = state['definition']
    prompt = PromptTemplate(
        template=prompt_for_reason(intro_and_defination=intro_and_defination),
        input_variables=[],
    )
    llm = model_load(alfaiz_api_key, temperature=0.7)
    reason = prompt | llm | StrOutputParser()
    reason = reason.invoke({})
    if reason != 'Not possible':
        state['reason'] = reason
    return state


def create_example_node(state):
    complete_string, state = complete_string_for_example(state)
    prompt = PromptTemplate(
        template=prompt_for_example(complete_answer=complete_string),
        input_variables=[],
    )
    llm = model_load(abhay_api_key, temperature=1)
    example = prompt | llm | StrOutputParser()
    example = example.invoke({})
    state['example'] = example
    return state


def create_real_life_example_node(state):
    complete_string, state = complete_string_for_example(state)
    prompt = PromptTemplate(
        template=prompt_for_real_life_example(complete_answer=complete_string),
        input_variables=[],
    )
    llm = model_load(asmit_api_key, temperature=1)
    real_example = prompt | llm | StrOutputParser()
    real_example = real_example.invoke({})
    state['real_life_example'] = real_example
    time.sleep(time_to_stop)
    return state


def complete_string_for_example(state):
    intro_and_defination = state['definition']
    reason = state['reason']
    types = state['types']
    difference = state['difference']
    advantage_and_disadvantage = state['advantage_and_disadvantage']

    complete_string = intro_and_defination

    if types is not None:
        complete_string += f"""\n {types}"""
    if difference is not None:
        complete_string += f"""\n {difference}"""
    if advantage_and_disadvantage is not None:
        complete_string += f"""\n {advantage_and_disadvantage}"""
    if reason is not None:
        complete_string += f"""\n {reason}"""

    return complete_string, state


def make_state_none(state):
    state['current_question'] = None
    state['example'] = None
    state['real_life_example'] = None
    state['advantage_and_disadvantage'] = None
    state['definition'] = None
    state['difference'] = None
    state['types'] = None
    state['reason'] = None

    return state


def complete_string_with_example(complete_string, state):
    example = state['example']
    real_life_example = state['real_life_example']
    complete_string += f"""\n {example} \n\n {real_life_example} \n"""
    return complete_string, state


def append_notes_in_docx_node(state):
    complete_string, state = complete_string_for_example(state)
    complete_string, state = complete_string_with_example(complete_string, state)
    append_notes_in_docx('final_output.docx', complete_string)
    state = make_state_none(state)
    return state


def create_new_docx(file_path):
    return Document() if not os.path.exists(file_path) else None


def add_text_to_new_page(doc, text, add_page_break):
    if add_page_break:
        doc.add_page_break()
    add_bold_paragraph(doc, text)
    return doc


def is_doc_empty(doc):
    return len(doc.paragraphs) == 0 or all(not p.text.strip() for p in doc.paragraphs)


def append_notes_in_docx(file_path, complete_string):
    doc = create_new_docx(file_path)
    if doc is None:
        doc = Document(file_path)
        add_page_break = not is_doc_empty(doc)
        doc = add_text_to_new_page(doc, complete_string, add_page_break)
    else:
        add_bold_paragraph(doc, complete_string)
    doc.save(file_path)


def add_bold_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    pattern = r'\*\*(.*?)\*\*'
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before the bold part
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        # Add bold text
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        last_end = match.end()
    # Add remaining text after the last bold part
    if last_end < len(text):
        paragraph.add_run(text[last_end:])